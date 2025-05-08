import docx
import os
import re

def extract_docx_to_markdown(docx_path, output_path):
    try:
        # Open the docx file
        doc = docx.Document(docx_path)
        
        # Initialize markdown content
        md_content = "# NSX Security Framework\n\n"
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading by looking at style name
                style_name = para.style.name.lower()
                if 'heading' in style_name:
                    # Determine heading level by the number in the style name
                    level = 1  # Default level
                    match = re.search(r'heading (\d+)', style_name)
                    if match:
                        level = int(match.group(1))
                    
                    # Add appropriate markdown heading markers
                    md_content += "#" * level + " " + para.text + "\n\n"
                else:
                    md_content += para.text + "\n\n"
        
        # Process tables
        for table in doc.tables:
            md_content += "\n| "
            
            # Get headers from first row
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
                md_content += cell.text.strip() + " | "
            
            md_content += "\n| "
            
            # Add separator row
            for _ in headers:
                md_content += "--- | "
            
            md_content += "\n"
            
            # Add table data
            for i, row in enumerate(table.rows):
                if i == 0:  # Skip header row as we've already processed it
                    continue
                
                md_content += "| "
                for cell in row.cells:
                    md_content += cell.text.strip() + " | "
                md_content += "\n"
            
            md_content += "\n"
        
        # Write to markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"Extraction complete. File saved as {output_path}")
        return True
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(current_dir, "NSX_Security_Framework.docx")
    output_path = os.path.join(current_dir, "NSX_Security_Framework.md")
    
    extract_docx_to_markdown(docx_path, output_path) 